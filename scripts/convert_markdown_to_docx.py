#!/usr/bin/env python
"""
Convert one or more Markdown/text files into a DOCX document.

Usage examples:
  python scripts/convert_markdown_to_docx.py \
    --input documentation/forensic_audit_2026-03-04/FULL_FORENSIC_AUDIT.md \
    --output documentation/forensic_audit_2026-03-04/FULL_FORENSIC_AUDIT.docx

  python scripts/convert_markdown_to_docx.py \
    --input documentation/forensic_audit_2026-03-04/FULL_FORENSIC_AUDIT.md \
            documentation/forensic_audit_2026-03-04/ANNEX_A_FILE_LEVEL_LEDGER.md \
    --output documentation/forensic_audit_2026-03-04/FULL_FORENSIC_AUDIT_WITH_ANNEX.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.shared import Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
CODE_FENCE_RE = re.compile(r"^\s*```")


def _add_code_paragraph(document: Document, line: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _write_markdown_block(document: Document, text: str) -> None:
    in_code_block = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")

        if CODE_FENCE_RE.match(line):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            _add_code_paragraph(document, line)
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            heading_text = heading_match.group(2).strip()
            document.add_heading(heading_text, level=level)
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        numbered_match = NUMBERED_RE.match(line)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            continue

        document.add_paragraph(line)


def convert_markdown_files_to_docx(
    input_paths: Iterable[Path],
    output_path: Path,
    *,
    title: str | None = None,
) -> None:
    paths: List[Path] = list(input_paths)
    if not paths:
        raise ValueError("No input markdown files provided.")

    document = Document()

    if title:
        document.add_heading(title, level=0)

    for idx, path in enumerate(paths):
        if not path.exists():
            raise FileNotFoundError(f"Input markdown not found: {path}")
        if path.suffix.lower() not in {".md", ".txt"}:
            raise ValueError(f"Input file must be .md or .txt: {path}")

        if len(paths) > 1:
            document.add_page_break()
            document.add_heading(path.name, level=1)

        content = path.read_text(encoding="utf-8", errors="ignore")
        _write_markdown_block(document, content)

        if idx < len(paths) - 1:
            document.add_paragraph("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown file(s) to DOCX.")
    parser.add_argument(
        "--input",
        nargs="+",
        required=True,
        help="One or more .md/.txt files.",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output .docx path.",
    )
    parser.add_argument(
        "--title",
        default=None,
        help="Optional document title.",
    )
    args = parser.parse_args()

    input_paths = [Path(p) for p in args.input]
    output_path = Path(args.output)
    convert_markdown_files_to_docx(input_paths, output_path, title=args.title)
    print(f"Wrote DOCX: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
